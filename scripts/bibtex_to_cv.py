#!/usr/bin/env python
"""
Reformat Google Scholar bibtex to Word document for CV.
Ensure that book chapters are listed as such in Google Scholar
"""
import argparse
import bibtexparser
from bibtexparser.bparser import BibTexParser
from bibtexparser.customization import convert_to_unicode
import docx

__author__ = 'Pim Bongaerts'
__copyright__ = 'Copyright (C) 2023 Pim Bongaerts'
__license__ = 'GPL'

def reformat_name(author_name):
    """ Reformat author name to Lastname, Firstname Initials """
    if not author_name:
        return ''
    name_split = author_name.replace('.', ' ').replace(',', ' ').split()
    last_name = name_split[0]
    first_names = name_split[1:len(name_split)]
    if len(first_names) == 1 and first_names[0].upper() == first_names[0]:
        initials = first_names[0]
    elif len(first_names) > 1 and first_names[1].upper() == first_names[1]:
        initials = '{0}{1}'.format(first_names[0][0],
                                   ''.join(name.upper() for name in first_names[1:]))
    else:
        initials = ''.join(name[0].upper() for name in first_names)
    if initials[len(initials)-3:] == 'VAN':
        #print(author_name, '->', '{0} van {1}'.format(last_name, initials[0:len(initials) - 3]))
        return last_name, '{0} van {1}'.format(last_name, initials[0:len(initials) - 3])
    else:
        #print(author_name, '->', '{0} {1}'.format(last_name, initials))
        return  last_name, '{0} {1}'.format(last_name, initials)

def reformat_doi(doi):
    """ Reformat doi """
    return doi.replace('https://doi.org/', '')

def output_to_docx(paragraph, publication, focal_authors):
    # Reformat coauthors
    authors_list = publication.get('author', '').split(' and ')
    if authors_list:
         for count, author_name in enumerate(authors_list):
            last_name, author_name_formatted = reformat_name(author_name)
            if count != 0:
                paragraph.add_run(', ').bold = False
            if last_name in focal_authors.split(','):
                paragraph.add_run(author_name_formatted).bold = True
            else:
                paragraph.add_run(author_name_formatted).bold = False

    paragraph.add_run(' ({}) '.format(publication.get('year', ''))).bold = False
    paragraph.add_run('{}. '.format(publication.get('title', ''))).bold = False
    if publication.get('journal', ''):
        paragraph.add_run('{} '.format(publication.get('journal', ''))).italic = True
        paragraph.add_run('{}:'.format(publication.get('volume', ''))).italic = False
    elif publication.get('booktitle', ''):
        paragraph.add_run('{} '.format(publication.get('booktitle', ''))).italic = True
    pages_reformat = publication.get('pages', '').replace('--', '–')
    paragraph.add_run('{}.'.format(pages_reformat)).italic = False

def main(filename, focal_authors):
    mydoc = docx.Document()

    with open(filename) as bibtex_file:
        parser = BibTexParser()
        parser.customization = convert_to_unicode
        bib_database = bibtexparser.load(bibtex_file, parser = parser)

    # Journal articles
    paragraph = mydoc.add_heading('Articles').bold = True
    for output_year in range(2030, 2000, -1): # hacky way to force reverse sort
        for publication in bib_database.entries:
            if publication.get('year', '') == str(output_year) and publication.get('booktitle', '') == '':
                paragraph = mydoc.add_paragraph()
                output_to_docx(paragraph, publication, focal_authors)

    # Book chapters
    paragraph = mydoc.add_heading('Book chapters').bold = True
    for output_year in range(2030, 2000, -1): # hacky way to force reverse sort
        for publication in bib_database.entries:
            if publication.get('year', '') == str(output_year) and publication.get('booktitle', '') != '':
                paragraph = mydoc.add_paragraph()
                output_to_docx(paragraph, publication, focal_authors)

    mydoc.save('publications.docx')

if __name__ == '__main__':
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument('filename', metavar='filename',
                        help='bibtex filename')
    parser.add_argument('focal_authors', metavar='focal_authors',
                        help='authors to be highlighted')
    args = parser.parse_args()
    main(args.filename, args.focal_authors)